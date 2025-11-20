import time
import io
import zipfile
import logging
from docx import Document
import openpyxl
import re
from datetime import datetime
from openpyxl.cell.cell import Cell
import locale

import os
import tempfile
import pythoncom  # Necessario per il threading di Windows
import win32com.client as win32



def check_consistency(excel_file: str, docx_file: str) -> tuple[bool, str, set[str]]:

    """
    Estrae tutte le stringhe racchiuse tra '{{' e '}}' da un file Word.

    Args:
        docx_path: Il percorso del file .docx.

    Returns:
        Una lista di stringhe trovate tra le parentesi graffe.
    """

    error_msg = []

    logging.info("CONSISTENCY_CHECK: started")


    try:
        # Carica il documento
        document = Document(docx_file)
    except Exception as e:
        logging.warning(f"CHECK_1: Cannot read word file {e}")
        error_msg = "Impossibile aprire il file word."
        return (False, error_msg, set())

    # Espressione Regolare (Regex):
    # \{\{    -> Corrisponde a '{{' (le graffe sono caratteri speciali e vanno 'escapate' con \)
    # (.*?)   -> Cattura qualsiasi carattere (punto) zero o più volte (asterisco) in modo non avido (punto interrogativo). 
    #            Questo è il contenuto che vogliamo estrarre.
    # \}\}    -> Corrisponde a '}}'
    pattern = re.compile(r"\{\{(.*?)\}\}")
    
    placeholders = set() # Usiamo un set per evitare duplicati

    # 1. Itera sui paragrafi (testo principale)
    for paragraph in document.paragraphs:
        # Trova tutte le corrispondenze nel testo del paragrafo
        matches = pattern.findall(paragraph.text)
        for match in matches:
            # Pulisci gli spazi bianchi attorno alla parola e aggiungi al set
            placeholders.add(match.strip())

    # 2. Itera sulle tabelle (testo nascosto nelle celle)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # Ricerca in tutti i paragrafi all'interno della cella
                for paragraph in cell.paragraphs:
                    matches = pattern.findall(paragraph.text)
                    for match in matches:
                        placeholders.add(match.strip())

    logging.debug(f"Placeholder values found in template model: {placeholders}")

    # --- 1. Caricamento ---
    try:
        workbook = openpyxl.load_workbook(excel_file,data_only=True)
        df = workbook.active
    except Exception as e:
        logging.warning(f"CHECK_1: cannot open the excel file: {e}")
        error_msg = "Impossibile aprire il file excel."
        return (False, error_msg, placeholders)
    

    
    # Questo dizionario mapperà: {'NomeColonna': indice_colonna_base_0}
    headers_set = set()

    # Itera solo sulla prima riga (min_row=1, max_row=1)
    for row in df.iter_rows(min_row=1, max_row=1):
        for cell in row:
            # cell.column è 1-based (A=1, B=2)
            # Lo convertiamo in 0-based (A=0, B=1) per usarlo con i tuple
            if cell.value: # Ignora eventuali celle vuote nell'intestazione
                headers_set.add(cell.value)

    logging.debug(f"Excel colums found: {headers_set}")

    if not headers_set.issuperset(placeholders):
        logging.warning(f"CHECK_1: Found some word placeholder not present in excel list: {placeholders.difference(headers_set)}")
        error_msg= f"Alcuni campi del word non sono stati trovati nell'excel: {placeholders.difference(headers_set)}."
        return  (False, error_msg, placeholders)
    else:
        logging.info("CONSISTENCY_CHECK: Ok")
        return (True, "", placeholders)

def convert_cell(mia_cella: Cell) -> str:
    """
    Prende in input una cella di openpyxl e restituisce una
    rappresentazione stringa formattata.

    - Le valute sono formattate come xxx.xxx,xx
    - Le date sono formattate come gg/mm/aaaa
    - Le celle vuote restituiscono una stringa vuota
    """
    
    valore = mia_cella.value
    formato = mia_cella.number_format
    
    # 1. Cella VUOTA
    if valore is None:
        return ""

    # 2. Cella DATA (Requisito: gg/mm/aaaa)
    # openpyxl converte automaticamente le date Excel in oggetti 'datetime'
    if isinstance(valore, datetime):
        return valore.strftime('%d/%m/%Y')

    # 3. Cella NUMERICA (Intero o Float)
    if isinstance(valore, (int, float)):
        
        # 3a. È una VALUTA? (Requisito: xxx.xxx,xx)
        # Controlliamo cercando i simboli di valuta nel formato Excel
        if '€' in formato or '$' in formato or '£' in formato or 'Currency' in formato:
            
            # "%.2f" forza 2 decimali
            # "grouping=True" usa il separatore delle migliaia ('.')
            # Il separatore decimale (',') è automatico grazie al locale.setlocale()
            return locale.format_string("%.2f", valore, grouping=True)
        
        # 3b. È un altro numero (generico o intero)
        else:
            # Per i numeri generici, usiamo la conversione standard di Python
            # (che userà il '.' come separatore decimale, es. 10.5)
            return str(valore)

    # 4. Cella STRINGA (Testo)
    if isinstance(valore, str):
        return valore
        
    # 5. Cella BOOLEANA o altro (es. Errori Excel come #N/D)
    # Per tutto il resto (True, False, ecc.), restituisce la sua
    # rappresentazione stringa standard.
    return str(valore)

def check_integrity(excel_file, placehoders: set[str]):

    excel_file.seek(0)

    # return dict

    dict_to_return = {}
    dict_errors = {}
        
    # 1. Collecting columns name with their index
    # logging.debug(f"Collecting excel column names and their index.")


    workbook = openpyxl.load_workbook(excel_file,data_only=True)
    df = workbook.active

    # Questo dizionario mapperà: {'NomeColonna': indice_colonna_base_0}
    headers_names = {}


    # Itera solo sulla prima riga (min_row=1, max_row=1)
    for row in df.iter_rows(min_row=1, max_row=1):
        for cell in row:
            # cell.column è 1-based (A=1, B=2)
            # Lo convertiamo in 0-based (A=0, B=1) per usarlo con i tuple
            if cell.value and cell.value in placehoders: 
                headers_names[cell.value] = cell.column-1



    # 2. Ciclo su ogni riga del file Excel
    
    colums_types = {}
    for row in df.iter_rows(min_row=2):

        # Se la riga è vuota warning e skippo
        is_row_empty = all(cell.value is None for cell in row)
        if is_row_empty:
            dict_errors[row[0].row] = "Empy row."
            logging.warning(f"Found empty row on line {row[0].row}. Skipped.")
            continue

        # logging.debug(f"Retrieving values for row {row[0].row}")

        try:
            values_of_row = {}
            for placeholder in placehoders:
                if  len(row) <= headers_names[placeholder] or row[headers_names[placeholder]].value is None :
                    #logger.warning(f"No value  found for placeholder {placeholder} on row {row[0].row}. Skipped row.")
                    #break
                    raise Exception(f"No value  found for placeholder {placeholder} on row {row[0].row}.")
                cell_got = row[headers_names[placeholder]]
                if placeholder in colums_types and colums_types[placeholder] != cell_got.data_type:
                    #logger.warning(f"Mismatch data type for {placeholder} in row {row[0].row}: expected {colums_types[placeholder]} - found {cell_got.data_type}. Skipped row.")
                    raise Exception(f"Mismatch data type for {placeholder} in row {row[0].row}: expected {colums_types[placeholder]} - found {cell_got.data_type}.")
                    #break
                elif placeholder  not in colums_types:
                    logging.debug(f"New data type recorded for {placeholder}: {cell_got.data_type}")
                    colums_types[placeholder] = cell_got.data_type
                values_of_row[placeholder] = convert_cell(cell_got)                            
                logging.debug(f"Row {row[0].row}: inserted {placeholder} = {values_of_row[placeholder]}")
        except Exception as e:
            dict_errors[row[0].row] = e
            logging.warning(f"Error found in excel file. {e}")
            continue
        dict_to_return[row[0].row] = values_of_row


    return (dict_to_return, dict_errors)


def check_integrity_paceholder(excel_file):
    """
    Verifica righe 'strane' nell'Excel.
    Return: list
    - list: Lista di warning (es. "Riga 5 saltata"). Vuota se tutto perfetto.
    """
    logging.info("CHECK_2: Inizio controllo integrità Excel")
    
    # TODO: INCOLLA QUI LA TUA LOGICA
    # Ricordati di resettare il puntatore del file se lo hai letto prima!
    excel_file.seek(0) 
    
    warnings = []
    time.sleep(1)
    
    # Esempio simulato:
    # warnings = ["Riga 10: Valore nullo, verrà saltata", "Riga 40: Formato data errato"]
    
    if warnings:
        logging.warning(f"CHECK_2: Trovati {len(warnings)} warning.")
    else:
        logging.info("CHECK_2: Nessun problema rilevato.")
        
    return warnings
def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    Sostituisce il testo in un paragrafo, gestendo i runs multipli 
    e preservando la formattazione originale.
    """
    if old_text not in paragraph.text:
        return False
    

    # 1. Trova l'indice di inizio e fine del segnaposto nel testo completo
    full_text = paragraph.text
    start_index = full_text.find(old_text)
    
    if start_index == -1:
        return False

    end_index = start_index + len(old_text)
    
    # 2. Inizializza le variabili per tracciare i runs
    current_len = 0
    
    # 3. Itera sui runs esistenti per trovare quali sono interessati dalla sostituzione
    for run in paragraph.runs:
        run_start = current_len
        run_end = current_len + len(run.text)

        # Il run inizia prima o al punto in cui inizia il segnaposto
        if run_start <= start_index < run_end:
            # Questo è il run iniziale dove avviene la sostituzione
            
            # 3a. Estrai il testo PRIMA del segnaposto (che deve rimanere)
            text_before = run.text[:start_index - run_start]
            
            # 3b. Se il segnaposto finisce all'interno di questo run (Caso più comune e semplice):
            if end_index <= run_end:
                text_after = run.text[end_index - run_start:]
                
                # Sostituzione nello stesso run
                run.text = text_before + new_text + text_after
                return True 
            
            # 3c. Se il segnaposto continua nel run successivo:
            else:
                # Sostituisce il testo del run con 'testo prima' + 'nuovo valore'
                run.text = text_before + new_text
                # Il resto del segnaposto sarà eliminato nei runs successivi o nell'ultimo run
                
        # Il run è completamente all'interno del segnaposto (da eliminare)
        elif start_index < run_start and run_end <= end_index:
            run.text = "" # Rimuove il testo di questo run

        # Il run è l'ultimo run interessato e contiene testo rimanente dopo il segnaposto
        elif run_start < end_index < run_end:
            text_after = run.text[end_index - run_start:]
            run.text = text_after
            
        current_len = run_end

    return True # Sostituzione completata (anche se ha attraversato più runs)

# --- STEP 3: GENERAZIONE FINALE ---
def generate_final_zip(word_file, campo_nome_file, righe_excel, placeholders_set):
    """
    Esegue l'elaborazione finale e crea lo zip.
    """
    logging.info("FINAL: Generation of final files...")
    word_file.seek(0)
    word_bytes_template = word_file.read()
    
    
    
    zip_buffer = io.BytesIO()
    with tempfile.TemporaryDirectory() as tmp_dir_path:
        logging.info(f"TEMP: Working on directory {tmp_dir_path}")
        
        # 1. GENERAZIONE DOCX SU DISCO
        # (Word non può aprire file dalla RAM, servono file veri)
        
        clienti_simulati = [
            {"nome": "Mario Rossi", "file": "Mario_Rossi"},
            {"nome": "Luigi Bianchi", "file": "Luigi_Bianchi"},
            {"nome": "Anna Verdi", "file": "Anna_Verdi"},
        ]
        
        paths_docx = [] # Ci segniamo i percorsi dei file creati

        # 2. Ciclo su ogni riga di dati
        for (index, row) in righe_excel.items():

            buffer_template = io.BytesIO(word_bytes_template)
            doc = Document(buffer_template)
          
            substitutions = {}

            try:
                for placeholder in placeholders_set:
                    substitutions[f"{{{{{placeholder}}}}}"] = row[placeholder]
            except Exception as e:
                logging.warning(f"Error in mapping substitutions: {e}. Row {index} skipped.")
                continue
            
            # Paragrafi
            for paragraph in doc.paragraphs:
                for old, new in substitutions.items():
                    found = replace_text_in_paragraph(paragraph, old, new)
                    if found == True:
                        logging.debug(f"Row {index}: substituting {old} -> {new} in paragraph.")

            
            # Campi tabella
            for table in doc.tables:
                for row_table in table.rows:
                    for cell in row_table.cells:
                        for paragraph in cell.paragraphs:
                            for old, new in substitutions.items():
                                found = replace_text_in_paragraph(paragraph, old, new) 
                                if found == True:
                                    logging.debug(f"Row {index}: substituting {old} -> {new} in table.")

            # 3. Salva il documento personalizzato
                    
            # Salva DOCX temporaneo
            try:
                filename_docx = f"{row[campo_nome_file]}.docx"
                path_docx = os.path.join(tmp_dir_path, filename_docx)
                doc.save(path_docx)
                logging.debug(f"File saved: {path_docx}")
                paths_docx.append(path_docx)
            except Exception as e:
                logging.warning(f"Error in creating file for row {index}: {e}. Skipped.")
                continue
        
            
        # 2. CONVERSIONE MASSIVA IN PDF (Win32)
        logging.info("PDF_CONVERSION: AvviStarting Word...")
        
        # FONDAMENTALE PER STREAMLIT: Inizializza COM nel thread corrente
        pythoncom.CoInitialize() 
        
        try:
            word_app = win32.Dispatch('Word.Application')
            word_app.Visible = False
            word_app.DisplayAlerts = False # Evita popup "Salvare modifiche?"
            
            for path_docx in paths_docx:
                try:
                    # Apri DOCX
                    doc_word = word_app.Documents.Open(path_docx)
                    
                    # Costruisci nome PDF
                    path_pdf = path_docx.replace(".docx", ".pdf")
                    
                    # Salva come PDF (FileFormat 17 = wdFormatPDF)
                    doc_word.SaveAs(path_pdf, FileFormat=17)
                    
                    doc_word.Close()
                    logging.info(f"PDF: Convertito {os.path.basename(path_pdf)}")
                    
                except Exception as e:
                    logging.warning(f"PDF ERROR: Cannot convert {path_docx}: {e}")
            
            word_app.Quit()
            
        except Exception as e:
            logging.warning(f"WIN32 ERROR: General error: {e}")
        finally:
            # Rilascia le risorse COM (importante per non bloccare il server)
            pythoncom.CoUninitialize()

        # 3. ZIPPARE I PDF
        logging.info("ZIP: Zipping files...")
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zf:
            for file_name in os.listdir(tmp_dir_path):
                full_path = os.path.join(tmp_dir_path, file_name)
                zf.write(full_path, file_name) # Salva solo il nome file, non il percorso

    logging.info("FINAL: Zip PDF completed.")
    return zip_buffer.getvalue()